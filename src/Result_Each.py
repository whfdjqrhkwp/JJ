#!/usr/bin/env python
# coding: utf-8

# In[1]:


import os
import sys
import re
import pandas as pd
import matplotlib.pyplot as pp
import seaborn as sb

#pp.switch_backend('agg')

from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

#from docx2pdf import convert
#import comtypes.client
import win32com.client

import import_ipynb

from Log import Logger
log = Logger('Result_Each')


# In[ ]:


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath('.')

    return os.path.join(base_path, relative_path)


# In[ ]:


class NAME(object):
    
    INFO_DIR = 'info_each'
    
    TEXT = {'morph_num' : 'morpheme_number.txt',
            'info_zero' : 'info_zero.txt',
            'info_UN' : 'info_UN.txt',
            'info_XR_zero' : 'info_XR_zero.txt',
            'scores_sentence' : 'scores_sentence.txt',
            'summary' : 'summary.txt'}
    
    IMAGE = {'score' : 'img_score.png',
             'score_voca' : 'img_score_voca.png',
             'score_gram' : 'img_score_gram.png',
             'hist' : 'img_hist.png',
             'hist_voca' : 'img_hist_voca.png',
             'hist_gram' : 'img_hist_gram.png',
             'rate_voca' : 'img_rate_voca.png',
             'rate_gram' : 'img_rate_gram.png'}
    


# In[2]:


class Information(object):
    
    @classmethod
    def __init__(cls, obj, address):
        cls.obj = obj
        cls.title = address.split('\\')[-1][:-4]
        cls.dir_address = get_result_address(address)
        
        log.info('Start Writting Information')
        cls.write_result_txt()
        cls.write_info_each()
        log.info('Complete Writting Information')
    
    @classmethod
    def write_result_txt(cls):
        #with open(os.path.join(cls.dir_address, 'info_all.txt'), 'w', encoding = 'utf-8') as file:
        #    file.write(cls.obj.get_buf_info())
        
        with open(os.path.join(cls.dir_address, NAME.TEXT['morph_num']), 'w', encoding = 'utf-8') as file:
            file.write(cls.obj.get_buf_morph_num())
        
        with open(os.path.join(cls.dir_address, NAME.TEXT['info_zero']), 'w', encoding = 'utf-8') as file:
            file.write(cls.obj.get_buf_info_zero())
        
        with open(os.path.join(cls.dir_address, NAME.TEXT['info_UN']), 'w', encoding = 'utf-8') as file:
            file.write(cls.obj.get_buf_info_UN())
            
        with open(os.path.join(cls.dir_address, NAME.TEXT['info_XR_zero']), 'w', encoding = 'utf-8') as file:
            file.write(cls.obj.get_buf_info_XR_zero())
        
        with open(os.path.join(cls.dir_address, NAME.TEXT['scores_sentence']), 'w', encoding = 'utf-8') as file:
            file.write(cls.obj.get_buf_scores_sentence())
        
        with open(os.path.join(cls.dir_address, NAME.TEXT['summary']), 'w', encoding = 'utf-8') as file:
            buf = 'File name : {}\n'.format(cls.title)
            buf += cls.obj.get_buf_summary()
            file.write(buf)
            
            
    @classmethod
    def write_info_each(cls):
        address = os.path.join(cls.dir_address, NAME.INFO_DIR)
        if not os.path.isdir(address):
            os.mkdir(address)
            
        info_each = cls.obj.info_each
        name = '{}.txt'
        delim = '\n{}\n'.format('-' * 50)
        
        zero_num = len(str(len(info_each)))
        for index, info in enumerate(info_each):
            index_string = str(index + 1).rjust(zero_num, '0')
            with open(os.path.join(address, name.format(index_string)), 'w', encoding = ' utf-8') as file:
                buf = delim.join(info)
                file.write(buf)
    
        


# In[6]:


class Graph(object):
    rate_label = ['{} points'.format(i + 1) for i in range(6)]
    rate_color = ['#90EE90', '#7FFFD4', '#FFD700', '#FF6347', '#EE82EE', '#323C73']
    rate_wedgeprops = {'width': 0.7, 'edgecolor': 'w', 'linewidth': 5}
    
    rwidth = 0.5
    
    isjupyter = True
    
    @classmethod
    def __init__(cls, obj, address):
        cls.obj = obj
        cls.dir_address = get_result_address(address)
        cls.bins = int(max(cls.obj.get_scores_sentence()) * 1.0) + 1
        
        log.info('Start Drawing Graph')
        
        cls.draw_score()
        cls.draw_hist()
        cls.draw_score_voca()
        cls.draw_hist_voca()
        cls.draw_score_gram()
        cls.draw_hist_gram()
        
        cls.draw_rate_voca()
        cls.draw_rate_gram()
        
        log.info('Complete Drawing Graph')
        
    @classmethod
    def draw_score(cls):
        df = pd.DataFrame(cls.obj.get_scores_sentence())
        
        fig = pp.figure(figsize = (12, 4))
        pp.title('Score of sentences')
        pp.plot(df)
        pp.xlabel('sentence number')
        pp.ylabel('score')
        pp.savefig(os.path.join(cls.dir_address, NAME.IMAGE['score']))
        
        if cls.isjupyter:
            pp.show()
            
        pp.close(fig)
        
    @classmethod
    def draw_score_voca(cls):
        df = pd.DataFrame(cls.obj.get_scores_voca_sentence())
        
        fig = pp.figure(figsize = (12, 4))
        pp.title('Vocabulary score of sentences')
        pp.plot(df)
        pp.xlabel('sentence number')
        pp.ylabel('vocabulary score')
        pp.savefig(os.path.join(cls.dir_address, NAME.IMAGE['score_voca']))
        
        if cls.isjupyter:
            pp.show()
            
        pp.close(fig)
        
    @classmethod
    def draw_score_gram(cls):
        df = pd.DataFrame(cls.obj.get_scores_gram_sentence())
        
        fig = pp.figure(figsize = (12, 4))
        pp.title('Grammar score of sentences')
        pp.plot(df)
        pp.xlabel('sentence number')
        pp.ylabel('grammar score')
        pp.savefig(os.path.join(cls.dir_address, NAME.IMAGE['score_gram']))
        
        if cls.isjupyter:
            pp.show()
            
        pp.close(fig)
        
    @classmethod
    def draw_hist(cls):
        df = pd.DataFrame(cls.obj.get_scores_sentence(), columns = ['score'])
        
        fig = pp.figure(figsize = (12, 4))
        pp.title('Hist of scores')
        pp.hist(df, bins = cls.bins, rwidth = cls.rwidth)
        pp.xlabel('score')
        pp.ylabel('number')
        pp.savefig(os.path.join(cls.dir_address, NAME.IMAGE['hist']))
        
        if cls.isjupyter:
            pp.show()
            
        pp.close(fig)
        
    @classmethod
    def draw_hist_voca(cls):
        df = pd.DataFrame(cls.obj.get_scores_voca_sentence(), columns = ['vocabulary score'])
        
        fig = pp.figure(figsize = (12, 4))
        pp.title('Hist of vocabulary scores')
        pp.hist(df, bins = cls.bins, rwidth = cls.rwidth)
        pp.xlabel('score')
        pp.ylabel('number')
        pp.savefig(os.path.join(cls.dir_address, NAME.IMAGE['hist_voca']))
        
        if cls.isjupyter:
            pp.show()
            
        pp.close(fig)
        
    @classmethod
    def draw_hist_gram(cls):
        df = pd.DataFrame(cls.obj.get_scores_gram_sentence(), columns = ['grammar score'])
        
        fig = pp.figure(figsize = (12, 4))
        pp.title('Hist of grammar scores')
        pp.hist(df, bins = cls.bins, rwidth = cls.rwidth)
        pp.xlabel('score')
        pp.ylabel('number')
        pp.savefig(os.path.join(cls.dir_address, NAME.IMAGE['hist_gram']))
        
        if cls.isjupyter:
            pp.show()
        
        pp.close(fig)
        
    @classmethod
    def draw_rate_voca(cls):
        voca_num = list(cls.obj.get_scores_voca_num())

        del voca_num[0]
    
        total_num = sum(voca_num)
        if total_num == 0:
            print('Vocabulary score ratio : All data are zero')
            return
 
        data = [round(num / total_num, 2) for num in voca_num]
       
        fig = pp.figure(figsize = (6, 6))
        pp.title('Vocabulary score ratio')
        pp.bar(range(len(data)), data, tick_label = cls.rate_label, color = cls.rate_color)
        #pp.pie(data, labels = cls.rate_label, autopct = '%.2f%%', colors = cls.rate_color , wedgeprops = cls.rate_wedgeprops, normalize = True)
        pp.savefig(os.path.join(cls.dir_address, NAME.IMAGE['rate_voca']))
        
        if cls.isjupyter:
            pp.show()
            
        pp.close(fig)

    @classmethod
    def draw_rate_gram(cls):
        gram_num = list(cls.obj.get_scores_gram_num())

        del gram_num[0]
     
        total_num = sum(gram_num)
        if total_num == 0:
            print('Grammar score ratio : All data are zero')
            return
 
        data = [round(num / total_num, 2) for num in gram_num]
    
        fig = pp.figure(figsize = (6, 6))
        pp.title('Grammar score ratio')
        pp.bar(range(len(data)), data, tick_label = cls.rate_label, color = cls.rate_color)
        #pp.pie(data, labels = cls.rate_label, autopct = '%.2f%%', colors = cls.rate_color , wedgeprops = cls.rate_wedgeprops, normalize = True)
        pp.savefig(os.path.join(cls.dir_address, NAME.IMAGE['rate_gram']))
        
        if cls.isjupyter:
            pp.show()
            
        pp.close(fig)


# In[4]:


class PDF(object):

    LINE_GRAPH_WIDTH = Cm(15.9)
    LINE_GRAPH_HEIGHT = Cm(5.29)
    BAR_GRAPH_WIDTH = Cm(9)
    BAR_GRAPH_HEIGHT = Cm(9)
    
    
    @classmethod
    def __init__(cls, obj, address):
        relative_form_address = os.path.join('Data', 'Format.docx')
        form_address = resource_path(relative_form_address)
        
        cls.document = Document(form_address)
        
        cls.obj = obj
        cls.dir_address = get_result_address(address)
        cls.title = address.split('\\')[-1][:-4]
        
        docx_address = os.path.join(cls.dir_address, '{}.docx'.format(cls.title))
        pdf_address = os.path.join(cls.dir_address, '{}.pdf'.format(cls.title))
        
        cls.fill_title()
        cls.fill_figure()
        cls.fill_table()
        
        cls.document.save(docx_address)
        
        try:
            cls.to_pdf(docx_address, pdf_address)
        except:
            return
        os.remove(docx_address)
        
    @classmethod
    def to_pdf(cls, docx_address, pdf_address):
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(docx_address)

        doc.SaveAs(pdf_address, FileFormat = 17)
        doc.Close()
        word.Quit()

    @classmethod
    def fill_title(cls):
        for p in cls.document.paragraphs:
            if '{filename}' in p.text:
                p.text = ''
                run = p.add_run(cls.title)
                run.bold = True
                run.font.size = Pt(12)
    
    @classmethod
    def fill_figure(cls):
        img_f1_1 = os.path.join(cls.dir_address, NAME.IMAGE['score'])
        img_f1_2 = os.path.join(cls.dir_address, NAME.IMAGE['hist'])
        img_f2_1 = os.path.join(cls.dir_address, NAME.IMAGE['score_voca'])
        img_f2_2 = os.path.join(cls.dir_address, NAME.IMAGE['hist_voca'])
        img_f3_1 = os.path.join(cls.dir_address, NAME.IMAGE['score_gram'])
        img_f3_2 = os.path.join(cls.dir_address, NAME.IMAGE['hist_gram'])
        
        img_f4 = os.path.join(cls.dir_address, NAME.IMAGE['rate_voca'])
        img_f5 = os.path.join(cls.dir_address, NAME.IMAGE['rate_gram'])
        
        
        for p in cls.document.paragraphs:
            if '{f1_1}' in p.text:
                p.text = ''
                p.add_run().add_picture(img_f1_1, width = cls.LINE_GRAPH_WIDTH, height = cls.LINE_GRAPH_HEIGHT)
            if '{f1_2}' in p.text:
                p.text = ''
                p.add_run().add_picture(img_f1_2, width = cls.LINE_GRAPH_WIDTH, height = cls.LINE_GRAPH_HEIGHT)
            if '{f2_1}' in p.text:
                p.text = ''
                p.add_run().add_picture(img_f2_1, width = cls.LINE_GRAPH_WIDTH, height = cls.LINE_GRAPH_HEIGHT)
            if '{f2_2}' in p.text:
                p.text = ''
                p.add_run().add_picture(img_f2_2, width = cls.LINE_GRAPH_WIDTH, height = cls.LINE_GRAPH_HEIGHT)
            if '{f3_1}' in p.text:
                p.text = ''
                p.add_run().add_picture(img_f3_1, width = cls.LINE_GRAPH_WIDTH, height = cls.LINE_GRAPH_HEIGHT)
            if '{f3_2}' in p.text:
                p.text = ''
                p.add_run().add_picture(img_f3_2, width = cls.LINE_GRAPH_WIDTH, height = cls.LINE_GRAPH_HEIGHT)
                
            if '{f4}' in p.text:
                p.text = ''
                p.add_run().add_picture(img_f4, width = cls.BAR_GRAPH_WIDTH, height = cls.BAR_GRAPH_HEIGHT)
            if '{f5}' in p.text:
                p.text = ''
                p.add_run().add_picture(img_f5, width = cls.BAR_GRAPH_WIDTH, height = cls.BAR_GRAPH_HEIGHT)
                
    
    @classmethod
    def fill_table(cls):
        tables = cls.document.tables
        
        cls.fill_table_1(tables[0])
        cls.fill_table_2(tables[1])
        cls.fill_table_3(tables[2])
        cls.fill_table_4(tables[3])
        cls.fill_table_5(tables[4])
    
    @classmethod
    def fill_cell(cls, cell, string):
        cell.text = ''
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = True
        run = cell.paragraphs[0].add_run(str(string))
        
    @classmethod
    def fill_table_1(cls, table):
        rows = table.rows
        
        cls.fill_cell(rows[1].cells[1], cls.obj.score)
        cls.fill_cell(rows[1].cells[2], max(cls.obj.scores_sentence))
        cls.fill_cell(rows[1].cells[3], min(cls.obj.scores_sentence))
        
        cls.fill_cell(rows[2].cells[1], len(cls.obj.sentences))
        cls.fill_cell(rows[2].cells[3], cls.obj.standard_deviation)
        
        cls.fill_cell(rows[3].cells[1], cls.obj.score_voca)
        cls.fill_cell(rows[3].cells[3], cls.obj.score_gram)
        
    @classmethod
    def fill_table_2(cls, table):
        rows = table.rows
        
        voca_num = list(cls.obj.get_scores_voca_num())
        gram_num = list(cls.obj.get_scores_gram_num())
        
        for index, cell in enumerate(rows[1].cells):
            if index > 0:
                cls.fill_cell(cell, voca_num[index])
        for index, cell in enumerate(rows[2].cells):
            if index > 0:
                cls.fill_cell(cell, gram_num[index])
                
    @classmethod
    def fill_table_3(cls, table):
        score_and_sentence = list(set(cls.obj.get_score_and_sentence()))
        score_and_sentence.sort(key = lambda score : -score[0])
        length = len(score_and_sentence)
        
        rows = table.rows
        row = rows[1]
        
        index = 0
        while True:
            cls.fill_cell(row.cells[0], score_and_sentence[index][0])
            cls.fill_cell(row.cells[1], score_and_sentence[index][1])
            
            if index + 1 < min(5, length):
                row = table.add_row()
                index += 1
            else:
                break
    
    @classmethod
    def fill_table_4(cls, table):
        score_and_sentence = list(set(cls.obj.get_score_and_sentence()))
        score_and_sentence.sort(key = lambda score : score[0])
        length = len(score_and_sentence)
        
        rows = table.rows
        row = rows[1]
        
        index = 0
        while True:
            cls.fill_cell(row.cells[0], score_and_sentence[index][0])
            cls.fill_cell(row.cells[1], score_and_sentence[index][1])
            
            if index + 1 < min(5, length):
                row = table.add_row()
                index += 1
            else:
                break
            
    @classmethod
    def fill_table_5(cls, table):
        morph_num = list(cls.obj.get_morph_num())
        
        morph_num_noun = list()
        morph_num_verb = list()
        morph_num_adjective = list()
        
        for e in morph_num:
            if 'NN' in e[0][1]:
                add = [e[0][0], e[1]]
                morph_num_noun.append(add)
            if 'VV' == e[0][1]:
                add = [e[0][0], e[1]]
                morph_num_verb.append(add)
            if 'VA' == e[0][1]:
                add = [e[0][0], e[1]]
                morph_num_adjective.append(add)
                
        morph_num_noun.sort(key = lambda num : -num[1])
        morph_num_verb.sort(key = lambda num : -num[1])
        morph_num_adjective.sort(key = lambda num : -num[1])
        
        len_noun = len(morph_num_noun)
        len_verb = len(morph_num_verb)
        len_adjective = len(morph_num_adjective)
        
        max_len = max(len_noun, len_verb, len_adjective)
        
        rows = table.rows
        row = rows[1]
        
        index = 0
        while True:
            if index < len_noun:
                cls.fill_cell(row.cells[0], morph_num_noun[index][0])
                cls.fill_cell(row.cells[1], morph_num_noun[index][1])

            if index < len_verb:
                cls.fill_cell(row.cells[2], morph_num_verb[index][0] + '다')
                cls.fill_cell(row.cells[3], morph_num_verb[index][1])

            if index < len_adjective:
                cls.fill_cell(row.cells[4], morph_num_adjective[index][0] + '다')
                cls.fill_cell(row.cells[5], morph_num_adjective[index][1])
            
            if index + 1 < min(34, max_len):
                row = table.add_row()
                index += 1
            else:
                break


# In[5]:


def get_result_address(address):
    
    def create_result_dir():
        address = os.path.join(os.getcwd(), 'Output')
        if not os.path.isdir(address):
            os.mkdir(address)
    
    def create_output_address(address, parent):
        if not parent:
            address = re.sub(pattern = r'\\[^\\]*.txt', repl = '', string = address)

        address = address.replace('.txt', '') 

        additional = address.replace(os.getcwd(), '')[1:]
        additional = additional.replace('Input', 'Output')

        result_address = os.path.join(os.getcwd(), additional)

        return result_address
    
    def create_output_dir(address):
        if not address.split('\\')[-2] == 'Input':
            addr = create_output_address(address, False)
            if not os.path.isdir(addr):
                os.mkdir(addr)

        dir_address = create_output_address(address, True)

        if not os.path.isdir(dir_address):
            os.mkdir(dir_address)

        return dir_address
    
    create_result_dir()
    address = create_output_dir(address)
    
    return address

